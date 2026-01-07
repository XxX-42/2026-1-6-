from docx import Document
import os

doc_path = r"D:\Documents\Codes\2025_Cuit_chinese-multimodal-emotion\2025_Cuit_chinese-multimodal-emotion_ppt&paper\2026-1-6 可行性分析\可行性分析报告模板(适用计科).doc"

try:
    doc = Document(doc_path)
    print("=== 文档内容 ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text}")
    
    print("\n=== 表格内容 ===")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- 表格 {t_idx} ---")
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            print(" | ".join(row_text))
except Exception as e:
    print(f"错误: {e}")
    print("尝试使用.docx格式...")
